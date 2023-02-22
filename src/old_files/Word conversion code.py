from html.parser import HTMLParser
from docx.enum.text import WD_COLOR_INDEX

# p = document.add_paragraph('', style='clovis-text')
p = document.add_paragraph("", style="clovis-text")
insertBorder(p, color="#888888")

mini_title = p.add_run("java".upper() + "\n", style="mini-title")
run = p.add_run("\n", style="after-mini-title")
mini_title.font.color.rgb = hex_color("#888888")


htmlStyle = {
    "built-in": 0,
    "keyword": 0,
    "number": 0,
    "string": 0,
}

stack = []


class MyHTMLParser(HTMLParser):
    def handle_starttag(self, tag, attrs):
        print("Encountered a start tag:", tag, attrs)
        attrs = dict(attrs)

        if tag == "span":
            if attrs["class"] == "hljs-built_in":
                htmlStyle["built-in"] = 1
                stack.append("hljs-built_in")
            elif attrs["class"] == "hljs-keyword":
                htmlStyle["keyword"] = 1
                stack.append("hljs-keyword")
            elif attrs["class"] == "hljs-number":
                htmlStyle["number"] = 1
                stack.append("hljs-number")
            elif attrs["class"] == "hljs-string":
                htmlStyle["string"] = 1
                stack.append("hljs-string")

    def handle_endtag(self, tag):
        print("Encountered an end tag :", tag)

        attrs = dict()
        if stack:
            attrs["class"] = stack.pop()

        if tag == "span":
            if attrs["class"] == "hljs-built_in":
                htmlStyle["built-in"] = 0
            elif attrs["class"] == "hljs-keyword":
                htmlStyle["keyword"] = 0
            elif attrs["class"] == "hljs-number":
                htmlStyle["number"] = 0
            elif attrs["class"] == "hljs-string":
                htmlStyle["string"] = 0

    def handle_data(self, data):
        print("Encountered some data  :", data)

        built_in_color = hex_color("#8757bd")
        keyword_color = hex_color("#446fbd")
        number_color = hex_color("#6d8600")
        string_color = hex_color("#e88501")

        run = p.add_run(data)
        run.font.name = "Courier New"

        if htmlStyle["built-in"]:
            run.font.color.rgb = built_in_color
        elif htmlStyle["keyword"]:
            run.font.color.rgb = keyword_color
        elif htmlStyle["number"]:
            run.font.color.rgb = number_color
        elif htmlStyle["string"]:
            run.font.color.rgb = string_color
        else:
            run.font.color.rgb = hex_color("#333333")


parser = MyHTMLParser()
"""parser.feed('<html><head><title>Test</title></head>'
            '<body><h1>Parse me!</h1></body></html>')
"""


a = '<div class="code-line"><span class="hljs-built_in">Set</span>&lt;<span class="hljs-built_in">String</span>&gt; ingredients = <span class="hljs-keyword">new</span> HashSet&lt;<span class="hljs-built_in">String</span>&gt;();</div><div class="code-line"></div><div class="code-line">ingredients.add(<span class="hljs-string">"chocolat"</span>);</div><div class="code-line">ingredients.remove(<span class="hljs-string">"chocolat"</span>);</div></div>'

# a = '<div class="code-line">N = <span class="hljs-keyword">int</span>(input())</div><div class="code-line"></div><div class="code-line">a = <span class="hljs-number">1</span></div><div class="code-line"></div><div class="code-line"><span class="hljs-keyword">for</span> i in range(N):</div><div class="code-line">&nbsp;&nbsp;&nbsp; a *= <span class="hljs-keyword">int</span>(input())%1000</div><div class="code-line"></div><div class="code-line">a = a%1000</div><div class="code-line"></div><div class="code-line"><span class="hljs-keyword">print</span>(<span class="hljs-string">"{:04d}"</span>.format(a))</div></div>'

b = a.split('<div class="code-line">')

for i in range(len(b)):
    if b[i] != "" and b[i][-6:] == "</div>":
        b[i] = b[i][:-6]

if b[0] == "":  # if first element is empty, delete it
    b = b[1:]


def add_line_number(line_number):
    """Ajoute le numéro de ligne line_number.
    (pour les blocs de code)"""
    padding_size = int(len(str(len(b))))  # number of digits

    l_nb = p.add_run(str(line_number).rjust(padding_size, " ") + " ")
    l_nb.font.name = "Courier New"
    l_nb.font.color.rgb = hex_color("#888888")


for i in range(len(b) - 1):  #
    add_line_number(i + 1)
    parser.feed(b[i])
    p.add_run("\n")

add_line_number(len(b))  # numérotation des lignes
parser.feed(b[-1])  # the last element (to not add a '\n' at the end)

# parser.feed(a)

### TEST background color
# Add a paragraph
# p = document.add_paragraph()

# Get the XML tag
tag = p._p

# Create XML element
shd = OxmlElement("w:shd")

# Add attributes to the element
# shd.set(qn('w:val'), 'clear')
shd.set(qn("w:color"), "auto")
shd.set(qn("w:fill"), "f8f8f8")

# This is the tricky part
tag.pPr.append(shd)

document.add_page_break()
document.save("demo2.docx")

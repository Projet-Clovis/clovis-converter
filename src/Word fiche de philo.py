## Import
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.style import WD_STYLE_TYPE
import os

os.chdir('C:\\Users\\AB2-PC\\Desktop')

## Fonctions
# ----- Insérer une ligne horizontale -----
# src : https://github.com/python-openxml/python-docx/issues/105
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn

def insertHR(paragraph):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '#dadce0') # on peut modifier la couleur ici
    pBdr.append(bottom)


def insertBorder(paragraph, color='#dadce0'):
    """ Add a border around the paragraph.

        * args
        - paragraph: the paragraph around which the border is set
        - color: the border color
        """
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color) # on peut modifier la couleur ici
    pBdr.append(bottom)

    top = OxmlElement('w:top')
    top.set(qn('w:val'), 'single')
    top.set(qn('w:sz'), '6')
    top.set(qn('w:space'), '1')
    top.set(qn('w:color'), color) # on peut modifier la couleur ici
    pBdr.append(top)

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '6')
    left.set(qn('w:space'), '1')
    left.set(qn('w:color'), color) # on peut modifier la couleur ici
    pBdr.append(left)

    right = OxmlElement('w:right')
    right.set(qn('w:val'), 'single')
    right.set(qn('w:sz'), '6')
    right.set(qn('w:space'), '1')
    right.set(qn('w:color'), color) # on peut modifier la couleur ici
    pBdr.append(right)


def hex_color(hex):
    """ Given a hex string (starting with '#' or without it),
        converts it to a RGBColor format (for python-docx).
    """
    assert (isinstance(hex, str) and len(hex) >= 6)

    if hex[0] == '#':
        hex = hex[1:]

    r, g, b = hex[0:2], hex[2:4], hex[4:6]
    r = '0x' + r
    g = '0x' + g
    b = '0x' + b

    r = int(r, base=16)
    g = int(g, base=16)
    b = int(b, base=16)

    return RGBColor(r, g, b)



## Ajouter un style personnalisé
# https://python-docx.readthedocs.io/en/latest/user/styles-using.html#add-or-delete-a-style


def creationStyle():
    """ Création des styles personnalisés. """
    styles = document.styles

    # ----- style h1 -----
    h1_color = '#bf3f3f'
    h1_style = styles.add_style('h1', WD_STYLE_TYPE.PARAGRAPH)
    h1_style.base_style = styles['Heading 1']
    h1_style.font.color.rgb = hex_color(h1_color)
    h1_style.font.name = 'Montserrat Medium'
    h1_style.font.size = Pt(20)

    # ----- style h2 -----
    h2_color = '#3ea546'
    h2_style = styles.add_style('h2', WD_STYLE_TYPE.PARAGRAPH)
    h2_style.base_style = styles['Heading 2']
    h2_style.font.color.rgb = hex_color(h2_color)
    h2_style.font.name = 'Montserrat Medium'
    h2_style.font.size = Pt(16)

    # ----- style h3 -----
    h3_color = '#6f9fd8'
    h3_style = styles.add_style('h3', WD_STYLE_TYPE.PARAGRAPH)
    h3_style.base_style = styles['Heading 3']
    h3_style.font.color.rgb = hex_color(h3_color)
    h3_style.font.name = 'Montserrat Medium'
    h3_style.font.size = Pt(14)

    # ----- style after-title -----
    # espace après les titres pour ne pas qu'ils se chevauchent
    after_title_style = styles.add_style('after-title', WD_STYLE_TYPE.PARAGRAPH)
    after_title_style.font.name = 'Montserrat Light'
    after_title_style.font.size = Pt(1)

    # ----- style text (clovis-text) -----
    text_color = '#333333'
    text_style = styles.add_style('clovis-text', WD_STYLE_TYPE.PARAGRAPH)
    text_style.base_style = styles['Normal']
    text_style.font.color.rgb = hex_color(text_color)
    text_style.font.name = 'Montserrat Light'
    text_style.font.size = Pt(12)

    # ----- style mini-title (colorful-block) -----
    mini_title_style = styles.add_style('mini-title', WD_STYLE_TYPE.CHARACTER)
    mini_title_style.base_style = styles['Normal']
    mini_title_style.font.name = 'Montserrat Light'
    mini_title_style.font.size = Pt(10)



## Liste des blocs
#todo : renvoyer le paragraphe à la fin de la fonction
def add_h1(text):
    """ Ajoute un titre h1. """
    insertHR(document.add_paragraph(text, style='h1'))
    document.add_paragraph('', style='after-title')


def add_h2(text):
    """ Ajoute un titre h2. """
    insertHR(document.add_paragraph(text, style='h2'))
    document.add_paragraph('', style='after-title')


def add_h3(text):
    """ Ajoute un titre h3. """
    insertHR(document.add_paragraph(text, style='h3'))
    document.add_paragraph('', style='after-title')


def add_text(text):
    """ Ajoute du texte. """
    document.add_paragraph(text, style='clovis-text')


def add_colorful_block(title, text, color='#4f4f4f'):
    """ Ajoute un bloc coloré. """
    p = document.add_paragraph('', style='clovis-text')
    insertBorder(p, color=color)

    title = p.add_run(title.upper() + '\n\n', style='mini-title')
    title.font.color.rgb = hex_color(color)

    p.add_run(text)


## Création du fichier
document = Document()

creationStyle()

add_h1('I - Le sujet')
add_h2('A) Problématiques')


add_text('Existe-t-il réellement un "moi-même", cette identité que je revendique fièrement ?')
add_text('Si elle existe, cette identité est-elle immuable ?')
add_text('Et est-elle nichée en moi ou n\'est-elle qu\'une apparence extérieure ?')
add_text('Définitions à connaître : éducation, alter ego, schizophrénie.')

add_h2('B) Le sujet selon Descartes')
add_text('Descartes donne un définition universelle du sujet, valable pour tous les humains. Le seul élément indiscutable auquel il parvient est qu\'il est sûr de penser. Une autre certitude émerge : s\'il pense, c\'est qu\'il existe. Donc si la pensée existe, l\'entité qui l\'exprime existe aussi. Même si elle s\'illusionne sur ce qui fait le réel. Pour Descartes, le sujet est donc à l\'origine des pensées. ce sujet derrière la pensée est appelé sujet pensant ou être pensant. La métaphysique de Descartes repose sur le cogito ergo sum, c\'est-à-dire, "Je pense donc je suis".')

add_h2('C) L\'inconscient théorisé par Freud')

add_text('Élaboré par Freud et Breuer à la fin du XIXè siècle, l\'hypothèse de l\'inconscient confirme que le sujet n\'est pas constitué que de la pensée consciente. Plusieurs forces agissent et influencent nos pensées et notre comportement.')
add_text('Freud affirme que le moi est multiple et soumis à des troubles de la personnalité : "Le moi n\'est pas maître dans sa propre maison". Il associe au "moi" un "ça" et un "surmoi" :')
add_text('- le moi est l\'individu conscient, capable de se plier aux conventions sociales;')
add_text('- le ça désigne les pulsions fortement liées au corps et à l\'instinct animal d\'un individu. Le plaisir domine le ça, qui est freiné par le moi;'
)
add_text('- le surmoi, enfin, est la morale, ce que la bonne éducation a inculqué à l\'enfant.')


p = document.add_paragraph('', style='clovis-text')
insertBorder(p, color='#ec9787')

g = p.add_run('RÉSUMÉ\n\n')
g.font.size = Pt(10)
g.font.color.rgb = hex_color('#ec9787')

p.add_run('Élaboré par Freud et Breuer à la fin du XIXè siècle, l\'hypothèse de l\'inconscient confirme que le sujet n\'est pas constitué que de la pensée consciente. Plusieurs forces agissent et influencent nos pensées et notre comportement.\n'
'Freud affirme que le moi est multiple et soumis à des troubles de la personnalité : "Le moi n\'est pas maître dans sa propre maison". Il associe au "moi" un "ça" et un "surmoi" :\n'
'- le moi est l\'individu conscient, capable de se plier aux conventions sociales;\n'
'- le ça désigne les pulsions fortement liées au corps et à l\'instinct animal d\'un individu. Le plaisir domine le ça, qui est freiné par le moi;\n'
'- le surmoi, enfin, est la morale, ce que la bonne éducation a inculqué à l\'enfant.')

#p.alignment = 3 # justified

add_colorful_block('test', 'hello world\nT\'es un ouf en fait j\'allais pas parler de soleil\nJe bosse chez Noz')

document.save('demo2.docx')

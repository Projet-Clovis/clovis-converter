from html.parser import HTMLParser
from docx.enum.text import WD_COLOR_INDEX

#p = document.add_paragraph('', style='clovis-text')
p = document.add_paragraph('', style='clovis-text')
insertBorder(p, color='#bd3d3a')

mini_title = p.add_run('attention'.upper() + '\n', style='mini-title')
p.add_run('\n', style='after-mini-title')
mini_title.font.color.rgb = hex_color('#bd3d3a')

htmlStyle = {
    'bold': 0,
    'italic': 0,
    'underline': 0,
    'sup': 0,
    'sub': 0,
    'f-code': 0,
}

class MyHTMLParser(HTMLParser):
    def handle_starttag(self, tag, attrs):
        print("Encountered a start tag:", tag, attrs)
        attrs = dict(attrs)

        if tag == 'b':
            htmlStyle['bold'] = 1
        elif tag == 'i':
            htmlStyle['italic'] = 1
        elif tag == 'u':
            htmlStyle['underline'] = 1
        elif tag == 'sup':
            htmlStyle['sup'] = 1
        elif tag == 'sub':
            htmlStyle['sub'] = 1
        elif tag == 'span':
            if attrs['class'] == 'f-code':
                htmlStyle['f-code'] = 1
            elif attrs['class'] == 'hl-yellow':
                htmlStyle['hl-yellow'] = 1


    def handle_endtag(self, tag):
        print("Encountered an end tag :", tag)

        if tag == 'b':
            htmlStyle['bold'] = 0
        elif tag == 'i':
            htmlStyle['italic'] = 0
        elif tag == 'u':
            htmlStyle['underline'] = 0
        elif tag == 'sup':
            htmlStyle['sup'] = 0
        elif tag == 'sub':
            htmlStyle['sub'] = 0
        elif tag == 'span':
            htmlStyle['f-code'] = 0
            #todo : mettre une pile pour savoir la classe correspondante ?

    def handle_data(self, data):
        print("Encountered some data  :", data)

        run = p.add_run(data)

        if htmlStyle['bold']:
            run.bold = True

        if htmlStyle['italic']:
            run.italic = True

        if htmlStyle['underline']:
            run.underline = True

        if htmlStyle['sup']:
            run.font.superscript = True

        if htmlStyle['sub']:
            run.font.subscript = True

        if htmlStyle['f-code']:
            run.font.name = 'Courier New'
            run.font.highlight_color = WD_COLOR_INDEX.GRAY_25



parser = MyHTMLParser()
"""parser.feed('<html><head><title>Test</title></head>'
            '<body><h1>Parse me!</h1></body></html>')
"""



a = 'On ne peut pas faire une liste d\'<span class="f-code">int</span> car les listes acceptent uniquement les <b>Objets </b>et pas les <b>types primitifs</b>.<br>'
a = 'On ne peut pas faire une liste d\'<span class="f-code">int</span> car les listes acceptent uniquement les <b>Objets </b>et pas les <b>types primitifs</b>.<br><br>Marco je ne rigole pas avec la fonction e<sup>x</sup> surtout avec x<sub>0</sub> = <u>lmao</u>.<br>J\'espère que <i><b>fusionner </b></i>les <span class="f-code"><i><u><b>st<sup>y</sup>l<sub>e</sub>s</b></u></i></span> fera pas tout bug.<br>'

if a[-4:] == '<br>': # if the 4 last characters is a <br>, remove it
    a = a[:-4] # without the 4 last characters

b = a.split('<br>')

for i in range(len(b)-1): #
    parser.feed(b[i])
    p.add_run('\n')
parser.feed(b[-1]) # the last element (to not add a '\n' at the end)

#parser.feed(a)



document.add_page_break()
document.save('demo2.docx')
